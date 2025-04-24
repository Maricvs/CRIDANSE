import os
import uuid
from typing import List, Optional, Dict, Any
from fastapi import UploadFile
from sqlalchemy.orm import Session
from sqlalchemy import func
from models.models import Document, DocumentChunk, Profile
from app.schemas.document_schema import DocumentCreate, DocumentChunkCreate, SearchQuery
import PyPDF2
import docx
import tiktoken
from openai import OpenAI
import numpy as np
import json
from datetime import datetime
from app.services.file_service import save_uploaded_file, delete_file, get_file_info

# Constants
UPLOAD_DIR = "uploads"
CHUNK_SIZE = 1000  # approximate number of words in a chunk
OVERLAP_SIZE = 200  # overlap between chunks
EMBEDDING_MODEL = "text-embedding-3-small"  # OpenAI model for embeddings

# Initialize OpenAI client
client = OpenAI()

# Function to count tokens
def num_tokens_from_string(string: str) -> int:
    """Returns the number of tokens in a string"""
    encoding = tiktoken.get_encoding("cl100k_base")
    num_tokens = len(encoding.encode(string))
    return num_tokens

# Function to create directories
def ensure_upload_dir():
    """Creates upload directory if it doesn't exist"""
    os.makedirs(UPLOAD_DIR, exist_ok=True)

# Function to get content from different document formats
def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extracts text from a file depending on its type"""
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type in ["docx", "doc"]:
        return extract_text_from_docx(file_path)
    elif file_type == "txt":
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

def extract_text_from_pdf(file_path: str) -> str:
    """Извлекает текст из PDF-файла"""
    text = ""
    with open(file_path, "rb") as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Извлекает текст из DOCX-файла"""
    doc = docx.Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def extract_text_from_txt(file_path: str) -> str:
    """Извлекает текст из TXT-файла"""
    with open(file_path, "r", encoding="utf-8") as file:
        text = file.read()
    return text

# Функция для разбиения текста на чанки
def split_text_into_chunks(text: str) -> List[str]:
    """Разбивает текст на чанки с перекрытием"""
    words = text.split()
    chunks = []
    
    for i in range(0, len(words), CHUNK_SIZE - OVERLAP_SIZE):
        chunk = " ".join(words[i:i + CHUNK_SIZE])
        if chunk:
            chunks.append(chunk)
    
    return chunks

# Функция для получения эмбеддингов
async def get_embedding(text: str) -> List[float]:
    """Получает эмбеддинг для текста с помощью API OpenAI"""
    response = client.embeddings.create(
        input=text,
        model=EMBEDDING_MODEL
    )
    return response.data[0].embedding

# Функция для сравнения эмбеддингов
def calculate_similarity(embedding1: List[float], embedding2: List[float]) -> float:
    """Рассчитывает косинусное сходство между двумя эмбеддингами"""
    dot_product = sum(a * b for a, b in zip(embedding1, embedding2))
    magnitude1 = sum(a * a for a in embedding1) ** 0.5
    magnitude2 = sum(b * b for b in embedding2) ** 0.5
    
    if magnitude1 * magnitude2 == 0:
        return 0
    
    return dot_product / (magnitude1 * magnitude2)

# Функции для работы с документами
async def upload_document(
    db: Session,
    user: Profile,
    file: UploadFile,
    title: str
) -> Document:
    """Загружает документ, сохраняет его и создает эмбеддинги"""
    try:
        # Сохраняем файл через единый сервис
        file_path, original_filename, file_size = await save_uploaded_file(file)
        
        # Создаем запись о документе в БД
        document_data = DocumentCreate(
            title=title,
            file_type=os.path.splitext(original_filename)[1].lower().lstrip('.')
        )
        
        db_document = Document(
            user_id=user.id,
            title=document_data.title,
            file_path=file_path,
            file_type=document_data.file_type,
            file_size=file_size,
            file_name=original_filename
        )
        
        db.add(db_document)
        db.commit()
        db.refresh(db_document)
        
        # Извлекаем текст из документа
        text = extract_text_from_file(file_path, document_data.file_type)
        
        # Разбиваем текст на чанки
        chunks = split_text_into_chunks(text)
        
        # Создаем эмбеддинги для каждого чанка
        for i, chunk_text in enumerate(chunks):
            embedding = await get_embedding(chunk_text)
            
            # Сохраняем чанк в БД
            chunk_data = DocumentChunkCreate(
                content=chunk_text,
                chunk_index=i,
                embedding=embedding
            )
            
            db_chunk = DocumentChunk(
                document_id=db_document.id,
                content=chunk_data.content,
                embedding=json.dumps(embedding),  # Сохраняем как JSON
                chunk_index=chunk_data.chunk_index
            )
            
            db.add(db_chunk)
        
        db.commit()
        return db_document
        
    except Exception as e:
        print(f"❌ [ERROR] Error uploading document: {str(e)}")
        # Если файл был сохранен, удаляем его
        if 'file_path' in locals():
            await delete_file(file_path)
        raise e

async def get_user_documents(
    db: Session,
    user: Profile
) -> List[Document]:
    """Получает список документов пользователя"""
    return db.query(Document).filter(Document.user_id == user.id).all()

async def delete_document(
    db: Session,
    document_id: int,
    user: Profile
) -> bool:
    """Удаляет документ и все его чанки"""
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == user.id
    ).first()
    
    if document:
        # Удаляем физический файл через единый сервис
        await delete_file(document.file_path)
        
        # Удаляем запись из БД
        db.delete(document)
        db.commit()
        return True
    
    return False

async def search_relevant_chunks(
    db: Session,
    user: Profile,
    search_query: SearchQuery
) -> List[Dict[str, Any]]:
    """Ищет релевантные чанки в документах пользователя"""
    # Получаем эмбеддинг запроса
    query_embedding = await get_embedding(search_query.query)
    
    # Получаем все документы пользователя
    documents = await get_user_documents(db, user)
    document_ids = [doc.id for doc in documents]
    
    if not document_ids:
        return []
    
    # Получаем все чанки для документов пользователя
    chunks = db.query(DocumentChunk).filter(
        DocumentChunk.document_id.in_(document_ids)
    ).all()
    
    # Сопоставляем каждый документ с его ID для быстрого доступа
    documents_by_id = {doc.id: doc for doc in documents}
    
    # Рассчитываем сходство для каждого чанка
    results = []
    for chunk in chunks:
        chunk_embedding = json.loads(chunk.embedding)  # Преобразуем из JSON
        similarity = calculate_similarity(query_embedding, chunk_embedding)
        
        results.append({
            "document_id": chunk.document_id,
            "document_title": documents_by_id[chunk.document_id].title,
            "chunk_id": chunk.id,
            "content": chunk.content,
            "similarity": similarity
        })
    
    # Сортируем по сходству и берем top-N
    results.sort(key=lambda x: x["similarity"], reverse=True)
    return results[:search_query.limit]

async def get_context_for_query(
    db: Session,
    user: Profile,
    query: str,
    limit: int = 5
) -> str:
    """Получает контекст для запроса из документов пользователя"""
    search_query = SearchQuery(query=query, limit=limit)
    relevant_chunks = await search_relevant_chunks(db, user, search_query)
    
    context = "\n\n---\n\n".join([chunk["content"] for chunk in relevant_chunks])
    return context

# Функция для генерации ответа с использованием GPT и контекста
async def generate_teacher_response(
    db: Session,
    user: Profile,
    query: str
) -> str:
    """Generates teacher's response based on document context"""
    # Get relevant context from documents
    context = await get_context_for_query(db, user, query)
    
    if not context:
        return "I apologize, but I couldn't find relevant information in the uploaded materials. Please upload study materials or clarify your question."

    # Form teacher prompt
    prompt = f"""You are a teacher who answers student questions based on provided study materials.
    
CONTEXT FROM STUDY MATERIALS:
{context}

STUDENT'S QUESTION:
{query}

Please answer the student's question using ONLY information from the provided context.
If there isn't enough information in the context for a complete answer, tell this to the student and answer based only on what is available in the context.
Be pedagogical, explain complex topics in simple language. Use examples from the context when necessary.
Always reference study materials, but don't mention the search process or embeddings in your answer."""
    
    # Call GPT to generate response
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are an experienced teacher who answers student questions based on study materials."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3,
        max_tokens=1500
    )
    
    return response.choices[0].message.content

# Функция для обработки содержимого документа
async def process_document_content(document_id: int, db: Session) -> bool:
    """
    Processes content of a document from "My Library":
    - extracts text
    - splits into chunks
    - creates embeddings
    - saves to database
    """
    print(f"🧠 [DEBUG] START: process_document_content({document_id})")
    
    # Get document from DB
    document = db.query(Document).filter(Document.id == document_id).first()
    
    if not document:
        raise ValueError(f"Document with ID {document_id} not found")
    
    # Extract text from document
    try:
        file_type = document.file_name.split('.')[-1] if document.file_name else document.file_type
        if not file_type:
            raise ValueError("Unknown file type")
            
        text = extract_text_from_file(document.file_path, file_type)
        
        # Split text into chunks
        chunks = split_text_into_chunks(text)
        
        print(f"✅ [INFO] Document {document.id} split into {len(chunks)} chunks")
        
        # Create embeddings for each chunk
        for i, chunk_text in enumerate(chunks):
            embedding = await get_embedding(chunk_text)
            
            # Save chunk to DB
            db_chunk = DocumentChunk(
                document_id=document.id,
                content=chunk_text,
                embedding=json.dumps(embedding),
                chunk_index=i
            )
            
            db.add(db_chunk)
        
        db.commit()
        print(f"✅ [INFO] Vectorization completed for document ID={document.id}")
        return True
    except Exception as e:
        print(f"❌ [ERROR] Error processing document {document_id}: {str(e)}")
        raise e